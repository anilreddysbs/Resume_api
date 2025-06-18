import docx
from docx2pdf import convert
import google.generativeai as genai
import re
import os

def get_section_text(doc, section_title):
    text = []
    capture = False
    for para in doc.paragraphs:
        if para.text.strip().lower() == section_title.lower():
            capture = True
            continue
        if capture and para.text.strip() and para.style.name.startswith('Heading'):
            break
        if capture:
            text.append(para.text)
    return "\n".join(text)

def replace_section_text_preserve_format(doc, section_title, new_text):
    capture = False
    start = None
    end = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == section_title.lower():
            capture = True
            start = i + 1
            continue
        if capture and para.text.strip() and para.style.name.startswith('Heading'):
            end = i
            break
    if start is None:
        print(f"Section '{section_title}' not found in the document.")
        return
    if end is None:
        end = len(doc.paragraphs)
    # Split new_text into lines
    new_lines = new_text.split('\n')
    para_idx = 0
    for i in range(start, end):
        if para_idx >= len(new_lines):
            break
        para = doc.paragraphs[i]
        # Replace only the text of the first run, keep formatting
        if para.runs:
            para.runs[0].text = new_lines[para_idx]
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.text = new_lines[para_idx]
        para_idx += 1
    # If there are more new lines than existing paragraphs, add them as new paragraphs
    for idx in range(para_idx, len(new_lines)):
        doc.add_paragraph(new_lines[idx])

def tailor_section_gemini(section_text, jd_text, api_key, section_title):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-flash')
    prompt = (
        f"Rewrite the following {section_title} section of a resume to better match this job description. "
        "Do NOT add or remove any bullet points, formatting, or structure. "
        "Keep the number of lines and the format exactly as in the original. "
        "Only update the wording to better fit the job description. "
        "Preserve the original formatting, bullet points, and structure. "
        "Ensure the rewritten section is concise, relevant, and tailored to the job description provided."
        "If you find 'Github Repository' word in the section,just use it as it is, do not change it.It has github repository links.\n\n"
        "Do not add extra bullets, headings, or sections.\n"
        f"Section:\n{section_text}\n\nJob Description:\n{jd_text}"
    )
    response = model.generate_content(prompt)
    return response.text

def extract_role_from_jd(jd_text):
    # Try to find 'Role:' or 'Position:' in the JD text
    match = re.search(r'(Role|Position)\s*:\s*([^\n\r]+)', jd_text, re.IGNORECASE)
    if match:
        role = match.group(2).strip()
        # Clean up for filename
        role = re.sub(r'[^a-zA-Z0-9]+', '_', role)
        return role
    return "UnknownRole"

if __name__ == "__main__":
    resume_path = "resume.docx"
    jd_path = "job_description.txt"
    GEMINI_API_KEY = "Gemini API Key"  # Replace with your Gemini API key

    doc = docx.Document(resume_path)
    jd_text = open(jd_path, encoding='utf-8').read()

    # Extract role for file naming
    role = extract_role_from_jd(jd_text)
    tailored_docx = f"tailored_resume_{role}.docx"
    tailored_pdf = f"tailored_resume_{role}.pdf"

    section_titles = [
        "PROFESSIONAL SUMMARY",
        "TECHNICAL SKILLS",
        "EDUCATION",
        "PROFESSIONAL EXPERIENCE",
        "PROJECTS",
        "PROFESSIONAL DEVELOPMENT",
        "ACHIEVEMENTS"
    ]

    for section_title in section_titles:
        section_text = get_section_text(doc, section_title)
        if not section_text.strip():
            print(f"No content found under section '{section_title}'. Skipping.")
            continue
        print(f"Tailoring section: {section_title}")
        tailored_text = tailor_section_gemini(section_text, jd_text, GEMINI_API_KEY, section_title)
        replace_section_text_preserve_format(doc, section_title, tailored_text)

    doc.save(tailored_docx)
    print(f"Tailored DOCX saved as {tailored_docx}")
    try:
        convert(tailored_docx, tailored_pdf)
        print(f"Tailored PDF saved as {tailored_pdf}")
    except Exception as e:
        print("PDF conversion failed:", e)